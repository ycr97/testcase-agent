# -*- coding: utf-8 -*-
"""Word 文档解析器（python-docx + LLM）"""

from __future__ import annotations

from pathlib import Path

import docx

from agent.parsers.text_parser import TextParser
from agent.parsers.base import BaseParser
from agent.schemas.api_schema import APIDocument


class WordParser(BaseParser):
    """解析 Word (.docx) 格式的接口文档。"""

    def __init__(self):
        self._text_parser = TextParser()

    def parse(self, source: str) -> APIDocument:
        """
        Args:
            source: Word 文件路径
        """
        path = Path(source)
        doc = docx.Document(str(path))

        # 提取所有段落和表格文本
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            parts.append("\n".join(rows))

        full_text = "\n\n".join(parts)
        api_doc = self._text_parser.parse(full_text)
        api_doc.title = api_doc.title or path.stem
        return api_doc
